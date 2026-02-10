from flask import Flask, render_template, request, send_file, jsonify
from docx import Document
from io import BytesIO

app = Flask(__name__)

@app.route('/')
def index():
    return render_template('index.html')


@app.route('/download', methods=['POST'])
def generate_document():
    data = request.get_json()

    event_name = data.get("eventName", "Unknown Event")
    event_location = data.get("eventLocation", "Unknown Location")
    start_date = data.get("startDate", "Unknown Start Date")
    end_date = data.get("endDate", "Unknown End Date")
    students = data.get("students", [])
    sender_name = data.get("senderName", "Unknown Student")

    # Optional field (not required but future-safe)
    send_email_option = data.get("sendEmailOption", False)

    if not students:
        return jsonify({"error": "No students selected"}), 400

    # 📄 Create Word document
    doc = Document()
    doc.add_paragraph("To\nClass Counsellor\nCSE Data Science\n\n")
    doc.add_paragraph(f"Subject: Request for Attendance Grant – {event_name} Participation\n\n")
    doc.add_paragraph(
        f"Respected Sir/Madam,\n\n"
        f"I hope you are doing well. I, {sender_name} from CSE Data Science, am participating in "
        f"{event_name} at {event_location} from {start_date} to {end_date}. "
        f"This event helps improve technical skills and teamwork.\n\n"
    )
    doc.add_paragraph("The participating members from our class are:\n")

    for student in students:
        doc.add_paragraph(student)

    doc.add_paragraph(
        "\nWe kindly request you to grant attendance for the mentioned dates.\n\n"
        "Looking forward to your kind consideration.\n\n"
        f"Sincerely,\n{sender_name}\nCSE Data Science"
    )

    # 🚀 Generate file in memory
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    return send_file(
        file_stream,
        as_attachment=True,
        download_name="Event_Attendance_Application.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


if __name__ == '__main__':
    app.run(debug=True)
